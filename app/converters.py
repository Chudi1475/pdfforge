"""PDF -> Word/Excel/HTML converters.

These are best-effort text-based extractions. Layout fidelity is limited
since PDF is a presentational format, but they're good enough for most
content reuse and editing.
"""
from __future__ import annotations
import html
import os
from pathlib import Path

import fitz
from PIL import Image


def pdf_to_docx(in_path: str, out_path: str, include_images: bool = True,
                progress=None) -> int:
    """Extract text + (optionally) images from each page into a .docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    src = fitz.open(in_path)
    docx = Document()
    page_count = src.page_count

    # default style
    style = docx.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    tmpdir = Path(out_path).parent / ("._" + Path(out_path).stem + "_tmp_images")
    tmpdir.mkdir(parents=True, exist_ok=True)

    try:
        for i in range(src.page_count):
            if i > 0:
                docx.add_page_break()
            page = src[i]
            d = page.get_text("dict")
            for block in d.get("blocks", []):
                if block.get("type", 0) == 0:  # text block
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        if not spans:
                            continue
                        para = docx.add_paragraph()
                        for span in spans:
                            run = para.add_run(span.get("text", ""))
                            run.font.size = Pt(max(6, int(span.get("size", 11))))
                            flags = span.get("flags", 0)
                            if flags & 16: run.bold = True
                            if flags & 2: run.italic = True
                            color_int = span.get("color", 0)
                            r = (color_int >> 16) & 0xff
                            g = (color_int >> 8) & 0xff
                            b = color_int & 0xff
                            if (r, g, b) != (0, 0, 0):
                                run.font.color.rgb = RGBColor(r, g, b)
                            fname = (span.get("font") or "").lower()
                            if "courier" in fname or "mono" in fname:
                                run.font.name = "Consolas"
                            elif "times" in fname or "serif" in fname:
                                run.font.name = "Times New Roman"
                elif include_images and block.get("type", 0) == 1:  # image block
                    try:
                        bbox = block.get("bbox", (0, 0, 0, 0))
                        w_pts = bbox[2] - bbox[0]
                        img_data = block.get("image")
                        if not img_data:
                            continue
                        fname = tmpdir / f"img_p{i}_b{block.get('number', 0)}.png"
                        with open(fname, "wb") as fout:
                            fout.write(img_data)
                        para = docx.add_paragraph()
                        run = para.add_run()
                        run.add_picture(str(fname), width=Inches(min(6.0, w_pts / 72)))
                    except Exception:
                        pass
            if progress:
                progress(i + 1, src.page_count)
        docx.save(out_path)
    finally:
        src.close()
        try:
            for f in tmpdir.iterdir():
                f.unlink()
            tmpdir.rmdir()
        except Exception:
            pass

    return page_count


def pdf_to_xlsx(in_path: str, out_path: str, progress=None) -> int:
    """Extract tables from each page into a workbook. Pages without tables
    get a plain text column on a fallback sheet."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    src = fitz.open(in_path)
    wb = Workbook()
    wb.remove(wb.active)
    total_tables = 0

    try:
        for i in range(src.page_count):
            page = src[i]
            tables = None
            try:
                tables = page.find_tables()
            except Exception:
                tables = None
            if tables and getattr(tables, "tables", None):
                for t_idx, tbl in enumerate(tables.tables):
                    try:
                        data = tbl.extract()
                    except Exception:
                        continue
                    if not data:
                        continue
                    ws = wb.create_sheet(title=f"P{i+1}T{t_idx+1}"[:31])
                    for r_idx, row in enumerate(data, start=1):
                        for c_idx, val in enumerate(row, start=1):
                            ws.cell(row=r_idx, column=c_idx, value=val)
                    # bold header
                    for cell in ws[1]:
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill("solid", fgColor="DDDDDD")
                    # autosize-ish
                    for col_cells in ws.columns:
                        col_letter = col_cells[0].column_letter
                        maxlen = max((len(str(c.value or "")) for c in col_cells), default=10)
                        ws.column_dimensions[col_letter].width = min(60, maxlen + 2)
                    total_tables += 1
            else:
                # fallback: dump text
                ws = wb.create_sheet(title=f"P{i+1}"[:31])
                lines = (page.get_text("text") or "").splitlines()
                for r_idx, line in enumerate(lines, start=1):
                    ws.cell(row=r_idx, column=1, value=line)
                ws.column_dimensions["A"].width = 100
            if progress: progress(i + 1, src.page_count)
        if not wb.sheetnames:
            wb.create_sheet("Empty")
        wb.save(out_path)
    finally:
        src.close()

    return total_tables


def pdf_to_html(in_path: str, out_path: str, embed_images: bool = True,
                progress=None) -> None:
    """Render each page's text + images to a single HTML document."""
    import base64
    src = fitz.open(in_path)
    title = Path(in_path).stem
    parts = [
        "<!doctype html>",
        f"<html><head><meta charset='utf-8'><title>{html.escape(title)}</title>",
        "<style>",
        "body{font-family:Calibri,Arial,sans-serif;max-width:900px;margin:24px auto;padding:0 24px;color:#222;}",
        ".page{padding:24px 0;border-bottom:1px dashed #bbb;}",
        ".page h2{color:#666;font-size:14px;font-weight:600;text-transform:uppercase;letter-spacing:1px;}",
        "img{max-width:100%;height:auto;display:block;margin:12px 0;}",
        "p{margin:6px 0;line-height:1.5;}",
        "</style></head><body>",
        f"<h1>{html.escape(title)}</h1>",
    ]
    try:
        for i in range(src.page_count):
            page = src[i]
            parts.append(f"<section class='page'><h2>Page {i+1}</h2>")
            d = page.get_text("dict")
            for block in d.get("blocks", []):
                if block.get("type", 0) == 0:
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        if not spans: continue
                        line_html = []
                        for span in spans:
                            txt = html.escape(span.get("text", ""))
                            flags = span.get("flags", 0)
                            size = int(span.get("size", 11))
                            style = f"font-size:{size}px;"
                            color_int = span.get("color", 0)
                            r = (color_int >> 16) & 0xff
                            g = (color_int >> 8) & 0xff
                            b = color_int & 0xff
                            if (r, g, b) != (0, 0, 0):
                                style += f"color:rgb({r},{g},{b});"
                            if flags & 16: txt = f"<strong>{txt}</strong>"
                            if flags & 2: txt = f"<em>{txt}</em>"
                            line_html.append(f"<span style='{style}'>{txt}</span>")
                        parts.append("<p>" + "".join(line_html) + "</p>")
                elif embed_images and block.get("type", 0) == 1:
                    try:
                        img_data = block.get("image")
                        if not img_data: continue
                        b64 = base64.b64encode(img_data).decode("ascii")
                        ext = block.get("ext", "png")
                        parts.append(f"<img src='data:image/{ext};base64,{b64}'/>")
                    except Exception:
                        pass
            parts.append("</section>")
            if progress: progress(i + 1, src.page_count)
        parts.append("</body></html>")
    finally:
        src.close()

    Path(out_path).write_text("\n".join(parts), encoding="utf-8")


def pdf_to_rtf(in_path: str, out_path: str):
    """Plain text wrapped in minimal RTF for compatibility."""
    src = fitz.open(in_path)
    body = []
    try:
        for i in range(src.page_count):
            if i > 0:
                body.append("\\page")
            txt = src[i].get_text("text")
            # escape RTF specials
            txt = txt.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
            txt = txt.replace("\n", "\\par ")
            body.append(txt)
    finally:
        src.close()
    rtf = "{\\rtf1\\ansi\\deff0 " + "".join(body) + "}"
    Path(out_path).write_text(rtf, encoding="utf-8")
